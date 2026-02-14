from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

# --- CONFIGURACIÓN DE COLORES (AJUSTA AQUÍ SI QUIERES OTRO TONO) ---
COLOR_FONDO_LATERAL = "2E4053"  # Un Gris Azulado Oscuro muy profesional (Hex sin #)
COLOR_TEXTO_LATERAL = RGBColor(255, 255, 255) # Blanco
COLOR_NOMBRE = RGBColor(46, 64, 83) # El mismo tono oscuro para el nombre
COLOR_SUBTITULOS = RGBColor(46, 64, 83) # Azul oscuro para títulos de sección

def set_cell_background(cell, color_hex):
    """Pinta el fondo de una celda usando XML"""
    shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color_hex))
    cell._tc.get_or_add_tcPr().append(shading_elm)

def crear_plantilla_final():
    doc = Document()
    
    # 1. MÁRGENES ESTRECHOS (Para que la barra lateral se vea completa)
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(1.27)
        section.bottom_margin = Cm(1.27)
        section.left_margin = Cm(1.27)
        section.right_margin = Cm(1.27)

    # 2. TABLA MAESTRA DE DISEÑO (Invisible)
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    
    # Anchos: 33% Izquierda, 67% Derecha
    table.columns[0].width = Inches(2.4)
    table.columns[1].width = Inches(5.0)

    # --- COLUMNA IZQUIERDA (BARRA LATERAL) ---
    cell_left = table.cell(0, 0)
    set_cell_background(cell_left, COLOR_FONDO_LATERAL)
    cell_left.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Función auxiliar para texto en la barra lateral
    def add_sidebar_line(text, size=10, bold=False, space_after=2):
        p = cell_left.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Segoe UI'
        run.font.size = Pt(size)
        run.font.color.rgb = COLOR_TEXTO_LATERAL
        run.bold = bold
        p.paragraph_format.space_after = Pt(space_after)
        return p

    # --- CONTENIDO BARRA LATERAL ---
    
    # Foto (Marcador)
    p_foto = add_sidebar_line("[ FOTO ]", size=14, bold=True, space_after=12)
    p_foto.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Contacto
    add_sidebar_line("CONTACTO", size=12, bold=True, space_after=6)
    add_sidebar_line("📍 {{ personal.direccion }}")
    add_sidebar_line("📱 {{ personal.telefono }}")
    add_sidebar_line("✉️ {{ personal.email }}")
    add_sidebar_line("_________________________", size=8) # Línea divisoria visual

    # Información Personal
    add_sidebar_line("\nINFORMACIÓN", size=12, bold=True, space_after=6)
    add_sidebar_line("CC: {{ personal.cedula }}")
    add_sidebar_line("Edad: {{ personal.edad }} Años")
    add_sidebar_line("Estado: {{ personal.estadoCivil }}")
    add_sidebar_line("Origen: {{ personal.lugarNacimiento }}")
    add_sidebar_line("_________________________", size=8)

    # Competencias (Skills)
    add_sidebar_line("\nCOMPETENCIAS", size=12, bold=True, space_after=6)
    # Jinja loop
    add_sidebar_line("{%p for skill in skills %}")
    add_sidebar_line("• {{ skill }}")
    add_sidebar_line("{%p endfor %}")
    add_sidebar_line("_________________________", size=8)

    # Diplomas
    add_sidebar_line("\nDIPLOMAS", size=12, bold=True, space_after=6)
    add_sidebar_line("{%p for item in diplomas %}")
    add_sidebar_line("🏆 {{ item }}", size=9)
    add_sidebar_line("{%p endfor %}")


    # --- COLUMNA DERECHA (CONTENIDO PRINCIPAL) ---
    cell_right = table.cell(0, 1)
    
    # Función auxiliar para texto principal
    def add_main_text(text, size=11, bold=False, color=None, style=None):
        p = cell_right.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Segoe UI'
        run.font.size = Pt(size)
        run.bold = bold
        if color: run.font.color.rgb = color
        return p

    # Nombre Gigante
    p_name = cell_right.add_paragraph()
    p_name.paragraph_format.space_before = Pt(10)
    run_name = p_name.add_run("{{ personal.nombre }}")
    run_name.font.name = 'Segoe UI Black' # Fuente gruesa
    run_name.font.size = Pt(26)
    run_name.font.color.rgb = COLOR_NOMBRE

    # Perfil Profesional
    add_main_text("PERFIL PROFESIONAL", size=12, bold=True, color=COLOR_SUBTITULOS).paragraph_format.space_before = Pt(12)
    p_perf = add_main_text("{{ personal.perfil }}")
    p_perf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Experiencia
    add_main_text("\nEXPERIENCIA LABORAL", size=12, bold=True, color=COLOR_SUBTITULOS).paragraph_format.space_before = Pt(12)
    # Borde inferior para el título
    
    # Loop Experiencia
    cell_right.add_paragraph("{%p for item in experiencia %}")
    
    # Cargo y Empresa
    p_exp_header = cell_right.add_paragraph()
    p_exp_header.paragraph_format.space_after = Pt(2)
    r_cargo = p_exp_header.add_run("{{ item.cargo }}")
    r_cargo.bold = True
    r_cargo.font.size = Pt(12)
    r_cargo.font.color.rgb = RGBColor(0, 0, 0)
    
    r_empresa = p_exp_header.add_run(" | {{ item.empresa }}")
    r_empresa.italic = True
    r_empresa.font.color.rgb = RGBColor(100, 100, 100) # Gris

    # Fecha
    p_date = cell_right.add_paragraph()
    p_date.paragraph_format.space_after = Pt(2)
    r_date = p_date.add_run("📅 {{ item.fecha }}")
    r_date.font.size = Pt(9)
    r_date.font.color.rgb = RGBColor(100, 100, 100)

    # Descripción
    p_desc = cell_right.add_paragraph()
    p_desc.add_run("{{ item.descripcion }}")
    p_desc.paragraph_format.space_after = Pt(12)
    
    cell_right.add_paragraph("{%p endfor %}")


    # Formación
    add_main_text("FORMACIÓN ACADÉMICA", size=12, bold=True, color=COLOR_SUBTITULOS)
    
    cell_right.add_paragraph("{%p for item in formacion %}")
    p_edu = cell_right.add_paragraph()
    p_edu.paragraph_format.space_after = Pt(0)
    r_edu = p_edu.add_run("🎓 {{ item.titulo }}")
    r_edu.bold = True
    
    p_edu_det = cell_right.add_paragraph()
    p_edu_det.add_run("{{ item.institucion }} - {{ item.fecha }}")
    p_edu_det.paragraph_format.space_after = Pt(8)
    cell_right.add_paragraph("{%p endfor %}")

    # GUARDAR
    import os
    if not os.path.exists("plantillas"):
        os.makedirs("plantillas")
        
    doc.save("plantillas/plantilla_cv.docx")
    print("✅ Plantilla 'Igualita' generada en: plantillas/plantilla_cv.docx")

if __name__ == "__main__":
    crear_plantilla_final()