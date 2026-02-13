from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def crear_plantilla_base():
    doc = Document()

    # 1. Crear una tabla invisible para la maquetación (2 columnas)
    # Columna 1: Barra lateral (Datos, Contacto, Skills)
    # Columna 2: Contenido principal (Perfil, Experiencia, Estudios)
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    
    # Ajustar anchos (aprox: 30% izquierda, 70% derecha)
    table.columns[0].width = Inches(2.0)
    table.columns[1].width = Inches(4.5)

    # --- COLUMNA IZQUIERDA (Datos Personales) ---
    cell_left = table.cell(0, 0)
    
    # Foto (Marcador de posición)
    p = cell_left.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("[FOTO AQUÍ]")
    run.bold = True
    
    # Datos de Contacto
    cell_left.add_paragraph("\nCONTACTO").runs[0].bold = True
    cell_left.add_paragraph("{{ personal.direccion }}")
    cell_left.add_paragraph("{{ personal.telefono }}")
    cell_left.add_paragraph("{{ personal.email }}")
    
    # Datos Demográficos
    cell_left.add_paragraph("\nDATOS PERSONALES").runs[0].bold = True
    cell_left.add_paragraph("Edad: {{ personal.edad }}")
    cell_left.add_paragraph("CC: {{ personal.cedula }}")
    cell_left.add_paragraph("Estado Civil: {{ personal.estadoCivil }}")
    cell_left.add_paragraph("Origen: {{ personal.lugarNacimiento }}")

    # Competencias (Loop)
    cell_left.add_paragraph("\nCOMPETENCIAS").runs[0].bold = True
    cell_left.add_paragraph("{% for skill in skills %}")
    cell_left.add_paragraph("• {{ skill }}")
    cell_left.add_paragraph("{% endfor %}")

    # --- COLUMNA DERECHA (Principal) ---
    cell_right = table.cell(0, 1)

    # Nombre Gigante
    p_nombre = cell_right.add_paragraph()
    run_nombre = p_nombre.add_run("{{ personal.nombre }}")
    run_nombre.bold = True
    run_nombre.font.size = Pt(24)
    run_nombre.font.color.rgb = RGBColor(0, 50, 100) # Azul oscuro

    # Perfil Profesional
    cell_right.add_paragraph("\nPERFIL PROFESIONAL").runs[0].bold = True
    cell_right.add_paragraph("{{ personal.perfil }}")

    # Experiencia Laboral (Loop)
    cell_right.add_paragraph("\nEXPERIENCIA LABORAL").runs[0].bold = True
    
    # Inicio del loop de experiencia
    cell_right.add_paragraph("{% for item in experiencia %}")
    
    p_cargo = cell_right.add_paragraph()
    run_cargo = p_cargo.add_run("{{ item.cargo }}")
    run_cargo.bold = True
    run_cargo.font.size = Pt(12)
    
    cell_right.add_paragraph("Empresa: {{ item.empresa }} | Fecha: {{ item.fecha }}")
    cell_right.add_paragraph("{{ item.descripcion }}")
    cell_right.add_paragraph("") # Espacio vacío
    cell_right.add_paragraph("{% endfor %}")

    # Formación Académica (Loop)
    cell_right.add_paragraph("\nFORMACIÓN ACADÉMICA").runs[0].bold = True
    cell_right.add_paragraph("{% for item in formacion %}")
    p_edu = cell_right.add_paragraph()
    p_edu.add_run("{{ item.titulo }}").bold = True
    cell_right.add_paragraph("{{ item.institucion }} - {{ item.fecha }}")
    cell_right.add_paragraph("{% endfor %}")
    
    # Diplomas (Loop)
    cell_right.add_paragraph("\nDIPLOMAS").runs[0].bold = True
    cell_right.add_paragraph("{% for item in diplomas %}")
    cell_right.add_paragraph("• {{ item }}")
    cell_right.add_paragraph("{% endfor %}")

    # Guardar
    doc.save("plantilla_mitu.docx")
    print("¡Plantilla creada con éxito! Ahora reinicia tu servidor uvicorn.")

if __name__ == "__main__":
    crear_plantilla_base()